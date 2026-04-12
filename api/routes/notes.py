"""
笔记系统 — 录音转写后的笔记存储、列表、详情、导出
"""

import io
from datetime import datetime
from typing import Optional, List

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse, HTMLResponse
from pydantic import BaseModel, Field
from database import get_db
from user_auth import get_current_user

router = APIRouter(prefix="/notes")


class CreateNoteRequest(BaseModel):
    title: str = Field(..., min_length=1, max_length=200)
    transcript: str = Field(default="")
    summary: str = Field(default="")
    key_points: list[str] = Field(default_factory=list)
    markdown: Optional[str] = None  # 文档模式生成的完整 Markdown


class UpdateNoteRequest(BaseModel):
    title: Optional[str] = None
    summary: Optional[str] = None
    key_points: Optional[List[str]] = None


@router.get("")
async def list_notes(
    limit: int = Query(default=50, ge=1, le=500),
    current_user: dict = Depends(get_current_user),
):
    """获取用户的笔记列表。"""
    db = get_db()
    result = (
        db.table("notes")
        .select("id, title, summary, key_points, markdown, created_at, updated_at")
        .eq("user_id", current_user["sub"])
        .order("created_at", desc=True)
        .limit(limit)
        .execute()
    )
    return {"notes": result.data}


@router.get("/{note_id}")
async def get_note(
    note_id: str,
    current_user: dict = Depends(get_current_user),
):
    """获取笔记详情。"""
    db = get_db()
    result = (
        db.table("notes")
        .select("*")
        .eq("id", note_id)
        .eq("user_id", current_user["sub"])
        .execute()
    )
    if not result.data:
        raise HTTPException(status_code=404, detail="Note not found")
    return result.data[0]


@router.post("")
async def create_note(
    req: CreateNoteRequest,
    current_user: dict = Depends(get_current_user),
):
    """创建笔记。"""
    db = get_db()
    row: dict = {
        "user_id": current_user["sub"],
        "title": req.title,
        "transcript": req.transcript,
        "summary": req.summary,
        "key_points": req.key_points,
    }
    if req.markdown is not None:
        row["markdown"] = req.markdown
    result = db.table("notes").insert(row).execute()
    return result.data[0]


@router.patch("/{note_id}")
async def update_note(
    note_id: str,
    req: UpdateNoteRequest,
    current_user: dict = Depends(get_current_user),
):
    """更新笔记。"""
    db = get_db()
    updates = {k: v for k, v in req.model_dump().items() if v is not None}
    if not updates:
        raise HTTPException(status_code=400, detail="Nothing to update")

    result = (
        db.table("notes")
        .update(updates)
        .eq("id", note_id)
        .eq("user_id", current_user["sub"])
        .execute()
    )
    if not result.data:
        raise HTTPException(status_code=404, detail="Note not found")
    return result.data[0]


@router.delete("/{note_id}")
async def delete_note(
    note_id: str,
    current_user: dict = Depends(get_current_user),
):
    """删除笔记。"""
    db = get_db()
    db.table("notes").delete().eq("id", note_id).eq("user_id", current_user["sub"]).execute()
    return {"status": "deleted"}


# ── 导出辅助 ──────────────────────────────────────────────────

def _fetch_note_for_export(note_id: str, user_id: str) -> dict:
    db = get_db()
    result = (
        db.table("notes").select("*")
        .eq("id", note_id).eq("user_id", user_id)
        .execute()
    )
    if not result.data:
        raise HTTPException(status_code=404, detail="Note not found")
    return result.data[0]


def _note_as_plain_text(note: dict) -> str:
    """笔记 → 纯文本（TXT 导出 / DOCX 兜底）"""
    lines = [note.get("title", "Untitled"), ""]
    created = note.get("created_at", "")
    if created:
        lines += [created[:10], ""]
    if note.get("markdown"):
        lines.append(note["markdown"])
    else:
        if note.get("summary"):
            lines += ["Summary", "───────", note["summary"], ""]
        if note.get("key_points"):
            lines += ["Key Points", "──────────"]
            for kp in note["key_points"]:
                lines.append(f"• {kp}")
            lines.append("")
        if note.get("transcript"):
            lines += ["Transcript", "──────────", note["transcript"]]
    return "\n".join(lines)


# ── TXT 导出 ──────────────────────────────────────────────────

@router.get("/{note_id}/export/txt")
async def export_note_txt(
    note_id: str,
    current_user: dict = Depends(get_current_user),
):
    note = _fetch_note_for_export(note_id, current_user["sub"])
    content = _note_as_plain_text(note)
    filename = note.get("title", "note").replace("/", "_")[:80] + ".txt"
    return StreamingResponse(
        io.BytesIO(content.encode("utf-8")),
        media_type="text/plain; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename*=UTF-8\'\'{filename}'},
    )


# ── DOCX 导出 ─────────────────────────────────────────────────

@router.get("/{note_id}/export/docx")
async def export_note_docx(
    note_id: str,
    current_user: dict = Depends(get_current_user),
):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    note = _fetch_note_for_export(note_id, current_user["sub"])
    doc = Document()

    # 标题
    title_para = doc.add_heading(note.get("title", "Untitled"), level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 日期
    created = note.get("created_at", "")[:10]
    if created:
        date_para = doc.add_paragraph(created)
        date_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        date_para.runs[0].font.size = Pt(10)

    doc.add_paragraph()  # 空行

    if note.get("markdown"):
        # 文档模式：逐行解析 Markdown 写入 DOCX
        for line in note["markdown"].splitlines():
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- [ ] "):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run("☐ " + line[6:])
            elif line.startswith("- [x] ") or line.startswith("- [X] "):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run("☑ " + line[6:])
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.strip() == "":
                doc.add_paragraph()
            else:
                doc.add_paragraph(line)
    else:
        # 摘要模式
        if note.get("summary"):
            doc.add_heading("Summary", level=2)
            doc.add_paragraph(note["summary"])
            doc.add_paragraph()

        if note.get("key_points"):
            doc.add_heading("Key Points", level=2)
            for kp in note["key_points"]:
                doc.add_paragraph(kp, style="List Bullet")
            doc.add_paragraph()

        if note.get("transcript"):
            doc.add_heading("Transcript", level=2)
            doc.add_paragraph(note["transcript"])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = note.get("title", "note").replace("/", "_")[:80] + ".docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename*=UTF-8\'\'{filename}'},
    )


# ── HTML 打印页（"导出 PDF" 的实际实现：浏览器 Ctrl+P）────────

@router.get("/{note_id}/export/print", response_class=HTMLResponse)
async def export_note_print(
    note_id: str,
    current_user: dict = Depends(get_current_user),
):
    note = _fetch_note_for_export(note_id, current_user["sub"])
    title   = note.get("title", "Untitled")
    created = note.get("created_at", "")[:10]

    # 把 Markdown 内容或摘要转成安全 HTML
    import html as html_lib

    def md_to_simple_html(md: str) -> str:
        """极简 Markdown → HTML（足够应付笔记格式）"""
        import re
        lines_out = []
        for line in md.splitlines():
            if line.startswith("# "):
                lines_out.append(f"<h1>{html_lib.escape(line[2:])}</h1>")
            elif line.startswith("## "):
                lines_out.append(f"<h2>{html_lib.escape(line[3:])}</h2>")
            elif line.startswith("### "):
                lines_out.append(f"<h3>{html_lib.escape(line[4:])}</h3>")
            elif line.startswith("- [ ] "):
                lines_out.append(f"<li>☐ {html_lib.escape(line[6:])}</li>")
            elif line.startswith("- [x] ") or line.startswith("- [X] "):
                lines_out.append(f"<li>☑ {html_lib.escape(line[6:])}</li>")
            elif line.startswith("- ") or line.startswith("* "):
                lines_out.append(f"<li>{html_lib.escape(line[2:])}</li>")
            elif line.strip() == "":
                lines_out.append("<br>")
            else:
                # **bold**
                escaped = html_lib.escape(line)
                escaped = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", escaped)
                lines_out.append(f"<p>{escaped}</p>")
        return "\n".join(lines_out)

    if note.get("markdown"):
        body_html = md_to_simple_html(note["markdown"])
    else:
        parts = []
        if note.get("summary"):
            parts.append(f"<h2>Summary</h2><p>{html_lib.escape(note['summary'])}</p>")
        if note.get("key_points"):
            items = "".join(f"<li>{html_lib.escape(kp)}</li>" for kp in note["key_points"])
            parts.append(f"<h2>Key Points</h2><ul>{items}</ul>")
        if note.get("transcript"):
            parts.append(f"<h2>Transcript</h2><p>{html_lib.escape(note['transcript'])}</p>")
        body_html = "\n".join(parts)

    html_content = f"""<!DOCTYPE html>
<html lang="zh">
<head>
  <meta charset="UTF-8">
  <title>{html_lib.escape(title)}</title>
  <style>
    @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+SC:wght@400;700&display=swap');
    body {{
      font-family: 'Noto Sans SC', 'Noto Sans', Arial, sans-serif;
      max-width: 720px; margin: 40px auto; padding: 0 24px;
      color: #111; line-height: 1.7; font-size: 15px;
    }}
    h1 {{ font-size: 24px; margin-bottom: 4px; }}
    h2 {{ font-size: 18px; margin-top: 24px; border-bottom: 1px solid #eee; padding-bottom: 4px; }}
    h3 {{ font-size: 15px; }}
    .meta {{ color: #888; font-size: 13px; margin-bottom: 24px; }}
    li {{ margin: 4px 0; }}
    ul {{ padding-left: 20px; }}
    @media print {{
      .no-print {{ display: none !important; }}
      body {{ margin: 0; padding: 0 16px; }}
    }}
  </style>
</head>
<body>
  <div class="no-print" style="margin-bottom:24px;padding:12px 16px;background:#f5f5f5;border-radius:8px;display:flex;align-items:center;gap:12px;">
    <span style="font-size:13px;color:#555;">Press <strong>Ctrl+P</strong> (or <strong>⌘P</strong>) and choose <em>Save as PDF</em></span>
    <button onclick="window.print()" style="margin-left:auto;padding:6px 16px;background:#111;color:#fff;border:none;border-radius:6px;cursor:pointer;font-size:13px;">Print / Save PDF</button>
  </div>
  <h1>{html_lib.escape(title)}</h1>
  <p class="meta">{created}</p>
  {body_html}
</body>
</html>"""

    return HTMLResponse(content=html_content)
